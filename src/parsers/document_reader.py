import os
# Stop PaddleOCR from pinging servers on startup
os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
from datetime import datetime
from tqdm import tqdm
import textwrap
import multiprocessing
import shutil
from functools import partial
from docx import Document
from tabulate import tabulate
import subprocess
import logging
import re
from paddleocr import PaddleOCR
import tempfile

# Initialize PaddleOCR globally so it doesn't reload the AI model for every single image
logging.getLogger("ppocr").setLevel(logging.ERROR) # Mute PaddleOCR spam


# It will only initialize the first time an image is processed
_ocr_instance=None
def get_ocr():
    global _ocr_instance
    if _ocr_instance is None:
        _ocr_instance = PaddleOCR(use_angle_cls=True, lang='ro')
    return _ocr_instance
# terminal progression bar
def _process_wrapper(args):
    func, instance, path = args
    print(f"\n⏳ [{datetime.now().strftime('%H:%M:%S')}] Started: {os.path.basename(path)}")
    res = func(instance, path)
    print(f"✅ [{datetime.now().strftime('%H:%M:%S')}] Finished: {os.path.basename(path)}")
    return (path, res)
'''
Reader Class:
    1) parses a file type (see the file_handler at the bottom) 
    
    2) extracts content as  
    
        "file": file_path,
        "title": file_title,
        "text": content
    
    to be written using generators/md_builder.py
    
    3) transforms path into tasks in a list and 
       calls multiprocess
'''

def process_file(func, path):  # Pre-binding self to the handler
    result = func(path)
    return (path, result)

class Reader:
    def __init__(self, folder_path, attachment_dir=None):
        self.folder = folder_path
        self.attachment_dir = attachment_dir
        
        # Ensure attachment directory exists if provided
        if self.attachment_dir:
            os.makedirs(self.attachment_dir, exist_ok=True)
        
    
    def readtxt(self, file_path):
        try:
            file_title = os.path.basename(file_path).replace(".txt", "")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()
            
            return {"file": file_path,
                    "title":file_title,
                    "text": content
            }
        except Exception as x:
            return {"file": file_path, "error": str(x)}

    def readocx(self, file_path):
        try:
            file_title = os.path.basename(file_path).replace(".docx", "")
            document = Document(file_path)
            all_content = [] # Store everything sequentially here

            # Wrap paragraphs
            for para in document.paragraphs:
                cleaned_text = para.text.strip()
                if cleaned_text:
                    wrapped_text = textwrap.fill(cleaned_text, width=80)
                    all_content.append(wrapped_text)

            # Format tables and append them at the end
            for table in document.tables:
                table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                formatted_table = tabulate(table_data, headers="firstrow", tablefmt="grid")
                all_content.append(f"\n{formatted_table}\n")

            links = {rel_id: rel._target for rel_id, rel in document.part.rels.items() if "hyperlink" in rel.reltype}

            # append the links to the bottom of the document
            if links:
                all_content.append("\n**Extracted Links:**")
                for link in links.values():
                    all_content.append(f"- {link}")

            # Join everything into one string
            final_markdown = "\n\n".join(all_content)

            return {
                "file": file_path,
                "title": file_title,
                "text": final_markdown, 
            }
        except Exception as x:
            return {"file": file_path, "error": str(x)}
        
    def readpho(self, file_path):
        try:
            file_title = os.path.splitext(os.path.basename(file_path))[0]
            
            # Use PaddleOCR (MinerU's internal engine)
            result = get_ocr().ocr(file_path, cls=True)
            
            text_lines = []
            
            # result[0] contains the list of detected text boxes
            if result and result[0]: 
                for line in result[0]:
                    # PaddleOCR returns data in format: [coordinates, (text, confidence_score)]
                    text = line[1][0] 
                    # --- DIACRITIC POST-PROCESSING ---
                    text = text.replace(',s', 'ș').replace(',S', 'Ș')
                    text = text.replace(',t', 'ț').replace(',T', 'Ț')
                    text = text.replace('~s', 'ș').replace('~t', 'ț')
                    text = text.replace('˘a', 'ă').replace('˘A', 'Ă')
                    text = text.replace('a˘', 'ă').replace('A˘', 'Ă') # Just in case it flips them
                    text = text.replace('ıˆ', 'î').replace('ˆı', 'î')
                    text = text.replace('Iˆ', 'Î').replace('ˆI', 'Î')
                    text = text.replace('ˆa', 'â').replace('ˆA', 'Â')
                    text = text.replace('aˆ', 'â').replace('Aˆ', 'Â')
                    text_lines.append(text)
                    
            final_markdown = "\n\n".join(text_lines)
            
            return {
                "file": file_path,
                "title": file_title,
                "text": final_markdown
            }

        except Exception as e:
            return {"file": file_path, "error": str(e)}
        
    #
    def readpdf(self, file_path):
        try:
            file_title = os.path.basename(file_path).replace(".pdf", "")
            temp_out_dir = os.path.join(tempfile.gettempdir(), f"mineru_temp_{file_title}")
            os.makedirs(temp_out_dir, exist_ok=True)

            try:
                # Run MinerU and capture the output
                mineru_process = subprocess.run(
                    ["magic-pdf", "-p", file_path, "-o", temp_out_dir, "-m", "auto"], 
                    check=True, 
                    capture_output=True,
                    text=True
                )
                
                # --- NEW: Print exactly what MinerU did ---
                print(f"\n[DEBUG] MinerU Output for {file_title}:")
                print(mineru_process.stdout)
                if mineru_process.stderr:
                    print(f"[DEBUG] MinerU Warnings/Errors: {mineru_process.stderr}")
                print("-" * 40)
                
            except subprocess.CalledProcessError as e:
                return {"file": file_path, "error": f"MinerU failed: {e.stderr}"}

            md_file_path = None
            images_dir = None
            
            for root, _, files in os.walk(temp_out_dir):
                for file in files:
                    if file.endswith('.md'):
                        md_file_path = os.path.join(root, file)
                        images_dir = os.path.join(root, "images")
                        break
                if md_file_path: 
                    break

            if not md_file_path or not os.path.exists(md_file_path):
                 # --- NEW: Do NOT delete the temp folder if it fails so we can inspect it manually ---
                 return {"file": file_path, "error": f"No .md file found. Go check the folder: {temp_out_dir}"}

            with open(md_file_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            if images_dir and os.path.exists(images_dir) and self.attachment_dir:
                for img_file in os.listdir(images_dir):
                    src_img_path = os.path.join(images_dir, img_file)
                    obsidian_img_name = f"{file_title}_{img_file}"
                    dest_img_path = os.path.join(self.attachment_dir, obsidian_img_name)
                    shutil.move(src_img_path, dest_img_path)
                    
                    old_img_tag = f"images/{img_file}"
                    md_content = md_content.replace(old_img_tag, obsidian_img_name)
            
            md_content = re.sub(r'!\[.*?\]\((.*?)\)', r'![[\1]]', md_content)

            # Cleanup only on success
            shutil.rmtree(temp_out_dir, ignore_errors=True)

            return {
                "file": file_path,
                "title": file_title,
                "text": md_content
            }

        except Exception as e:
            return {"file": file_path, "error": str(e)}
        
    # Convert .doc to .docx using antiword
    def readdoc(self, file_path):
        try:
            
            docx_path = file_path + "x"

            if not os.path.exists(docx_path):
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
                text = result.stdout

                doc = Document()
                doc.add_paragraph(f"Note: {file_path} is a legacy .doc file. Only text was extracted.\n")
                doc.add_paragraph(text)
                doc.save(docx_path)

            extracted_content = self.readocx(docx_path)
            os.remove(docx_path)
            return extracted_content

        except Exception as e:
            return {"file": file_path, "error": str(e)}

    # Scan for files recursevly (maintain order),
    # gets file paths, appends them to a list for multiprocces
    def scanner(self, folder,use_multiprocessing=True):  
        file_entries, sub_dirs, tasks = [], [], []
        results = []
        processed_files = set()

        with os.scandir(folder) as entries:
            for entry in entries:
                if entry.is_file() and entry.path not in processed_files:
                    file_entries.append(entry.path)
                    processed_files.add(entry.path)
                elif entry.is_dir():
                    sub_dirs.append(entry.path)  # Separate files from directories

        for file_path in file_entries:
            _, extension = os.path.splitext(file_path)
            handler = file_handlers.get(extension.lower())

            if handler:
                tasks.append((handler, self, file_path))
        
        try:
            if use_multiprocessing and tasks:
                # Default to 2 workers to prevent system crashes on heavy ML tasks.
                with multiprocessing.Pool(processes=2) as pool:
                    for result in tqdm(pool.imap_unordered(_process_wrapper, tasks), total=len(tasks), desc="Overall Progress", unit="file"):
                        results.append(result)
            else:
                for task in tasks:
                    results.append(process_file(*task))
        except Exception as e:
            results.append(("System_Error", {"error": f"Multiprocessing error: {e}"}))
        for sub_dir in sub_dirs:  # Recursively scan subdirectories
            results += self.scanner(sub_dir,use_multiprocessing=use_multiprocessing)

        return results

file_handlers = {
    ".txt": Reader.readtxt,
    ".docx": Reader.readocx,
    ".png": Reader.readpho,
    ".jpeg": Reader.readpho,
    ".jpg": Reader.readpho,
    ".webp": Reader.readpho,
    ".bmp": Reader.readpho,
    ".pdf":Reader.readpdf,
    ".doc":Reader.readdoc,
}  


#    [os.remove(e.path) for e in os.scandir(folder_path)]
#    file remover

            